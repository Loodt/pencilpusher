"""Document readers — convert source documents to LLM-consumable formats.

Primary reader: MarkItDown (Microsoft, 96K stars) converts any document to Markdown.
Fallback readers: PyMuPDF for PDF images, python-docx for DOCX text.
"""

from pathlib import Path


def read_with_markitdown(file_path: Path) -> str:
    """Convert any supported document to Markdown using Microsoft MarkItDown.

    This is the preferred reader — produces structured Markdown from PDF, DOCX,
    PPTX, XLSX, images, and more. The Markdown is then sent to Claude's text API
    (cheaper and faster than vision API).

    Returns empty string if MarkItDown fails or produces no useful content.
    """
    try:
        from markitdown import MarkItDown
        md = MarkItDown()
        result = md.convert(str(file_path))
        text = result.text_content.strip() if result.text_content else ""
        # Only return if we got meaningful content (not just whitespace)
        if len(text) > 20:
            return text
        return ""
    except Exception:
        return ""


def read_pdf_as_images(pdf_path: Path) -> list[tuple[bytes, str]]:
    """Convert each page of a PDF to a PNG image for Claude vision.

    Returns list of (image_bytes, media_type) tuples.
    """
    import fitz  # PyMuPDF

    doc = fitz.open(str(pdf_path))
    images = []

    for page in doc:
        # Render at 200 DPI for good quality without being huge
        mat = fitz.Matrix(200 / 72, 200 / 72)
        pix = page.get_pixmap(matrix=mat)
        images.append((pix.tobytes("png"), "image/png"))

    doc.close()
    return images


def read_pdf_as_text(pdf_path: Path) -> str:
    """Extract text from a PDF. Falls back to empty string if no text layer."""
    import fitz

    doc = fitz.open(str(pdf_path))
    text_parts = []

    for page in doc:
        text = page.get_text()
        if text.strip():
            text_parts.append(text)

    doc.close()
    return "\n\n---\n\n".join(text_parts)


def read_docx_as_text(docx_path: Path) -> str:
    """Read a Word document as plain text with structure preserved."""
    import docx

    doc = docx.Document(str(docx_path))
    parts = []

    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)

    # Also extract from tables
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts)


def read_image(image_path: Path) -> tuple[bytes, str]:
    """Read an image file, returning (bytes, media_type)."""
    suffix = image_path.suffix.lower()
    media_types = {
        ".png": "image/png",
        ".jpg": "image/jpeg",
        ".jpeg": "image/jpeg",
        ".gif": "image/gif",
        ".webp": "image/webp",
    }
    media_type = media_types.get(suffix, "image/png")
    return image_path.read_bytes(), media_type


def detect_file_type(path: Path) -> str:
    """Detect file type from extension."""
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return "pdf"
    elif suffix in (".docx", ".doc"):
        return "docx"
    elif suffix in (".png", ".jpg", ".jpeg", ".gif", ".webp", ".tiff", ".bmp"):
        return "image"
    else:
        return "unknown"
