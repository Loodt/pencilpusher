"""Word document filler — write matched values into DOCX files.

Three proven approaches (tested on real documents):

1. SDT content controls — zipfile + lxml at the XML level (bypasses python-docx limitations)
2. Table cells — python-docx table cell access (for government forms with table-based layouts)
3. Placeholder runs — python-docx run-level text replacement (for [___] and _____ patterns)

Key constraint: preserve ALL styling (fonts, sizes, colors, formatting).
"""

import re
import zipfile
from pathlib import Path

from lxml import etree

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NSMAP = {"w": W}


def fill_docx(
    docx_path: Path,
    matches: list[dict],
    fields: list,
    output_path: Path,
) -> Path:
    """Fill a Word document with matched values using all three approaches.

    1. First pass: fill SDT content controls via raw XML (most reliable)
    2. Second pass: fill table cells adjacent to label cells
    3. Third pass: replace placeholder patterns in paragraph runs
    """
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    # Build lookups
    fill_map = {}  # field_key -> value
    name_map = {}  # field_name -> value
    for match in matches:
        val = match.get("matched_value")
        if not val:
            continue
        if match.get("field_key"):
            fill_map[match["field_key"]] = val
        if match.get("field_name"):
            name_map[match["field_name"]] = val

    # Pass 1: SDT filling via zipfile + lxml (proven approach)
    sdt_filled = _fill_sdts_raw_xml(docx_path, output_path, fill_map)

    # Use the SDT-filled output as input for subsequent passes
    working_path = output_path if sdt_filled else docx_path

    # Pass 2 + 3: table cells and placeholders via python-docx
    import docx
    doc = docx.Document(str(working_path))

    # Pass 2: table cell filling
    for table in doc.tables:
        _fill_table_cells(table, fill_map, name_map)

    # Pass 3: placeholder replacement in paragraphs
    for para in doc.paragraphs:
        _replace_placeholders(para, fill_map, name_map)

    # Also check paragraphs inside table cells (some forms have placeholders in tables)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_placeholders(para, fill_map, name_map)

    doc.save(str(output_path))
    return output_path


def _fill_sdts_raw_xml(docx_path: Path, output_path: Path, fill_map: dict) -> bool:
    """Fill SDT content controls by directly editing document.xml inside the ZIP.

    This bypasses python-docx entirely — proven to work where python-docx fails
    (issue #965: SDT edits don't persist on save).

    Returns True if any SDTs were found and filled.
    """
    filled_count = 0

    with zipfile.ZipFile(str(docx_path), "r") as zin:
        with zipfile.ZipFile(str(output_path), "w", zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)

                if item.filename == "word/document.xml":
                    tree = etree.fromstring(data)

                    for sdt in tree.iter(f"{{{W}}}sdt"):
                        # Try tag first, then alias
                        tag_el = sdt.find(".//w:sdtPr/w:tag", NSMAP)
                        alias_el = sdt.find(".//w:sdtPr/w:alias", NSMAP)

                        key = None
                        if tag_el is not None:
                            key = tag_el.get(f"{{{W}}}val")
                        if key not in fill_map and alias_el is not None:
                            key = alias_el.get(f"{{{W}}}val")

                        if key and key in fill_map:
                            value = fill_map[key]
                            # Set text in first w:t, clear the rest
                            t_elements = sdt.findall(".//w:sdtContent//w:t", NSMAP)
                            if t_elements:
                                t_elements[0].text = value
                                for extra in t_elements[1:]:
                                    extra.text = ""
                                filled_count += 1

                    if filled_count > 0:
                        data = etree.tostring(
                            tree, xml_declaration=True, encoding="UTF-8", standalone=True
                        )

                zout.writestr(item, data)

    return filled_count > 0


def _fill_table_cells(table, fill_map: dict, name_map: dict) -> None:
    """Fill empty table cells adjacent to label cells.

    This handles government forms where labels are in left columns
    and data entry is in right columns (empty cells).
    """
    for row in table.rows:
        cells = row.cells
        for i, cell in enumerate(cells):
            label = cell.text.strip().rstrip(":")
            if not label:
                continue
            # Check if any subsequent cell in this row is empty (the data entry cell)
            for j in range(i + 1, len(cells)):
                target = cells[j]
                if not target.text.strip():
                    # Try to match the label to a value
                    value = name_map.get(label) or name_map.get(label + ":")
                    if value:
                        # Write into the first paragraph of the empty cell
                        if target.paragraphs:
                            from docx.shared import Pt
                            run = target.paragraphs[0].add_run(str(value))
                            run.font.size = Pt(10)
                    break  # Only fill the first empty cell after a label


def _replace_placeholders(para, fill_map: dict, name_map: dict) -> None:
    """Replace placeholder patterns in a paragraph, preserving run formatting."""
    full_text = para.text
    if not full_text:
        return

    # Literal match against fill_map keys
    for placeholder, value in fill_map.items():
        if placeholder in full_text:
            _replace_text_in_runs(para.runs, placeholder, value)
            return

    # Pattern-based matching
    patterns = [
        (r"\[___+\]", None),        # [_____]
        (r"\{(\w[\w\s]*)\}", 1),    # {field_name}
        (r"<<(\w[\w\s]*)>>", 1),    # <<field_name>>
        (r"_{5,}", None),           # _________
    ]

    for pattern, group in patterns:
        for match in re.finditer(pattern, full_text):
            matched_text = match.group(0)

            value = fill_map.get(matched_text)
            if not value and group and match.lastindex and match.lastindex >= group:
                field_name = match.group(group)
                value = fill_map.get(field_name) or name_map.get(field_name)

            # Context-based: use text before placeholder as label
            if not value:
                prefix = full_text[: match.start()].strip().rstrip(":")
                value = name_map.get(prefix)

            if value:
                _replace_text_in_runs(para.runs, matched_text, str(value))


def _replace_text_in_runs(runs, old_text: str, new_text: str) -> None:
    """Replace text across runs while preserving each run's formatting."""
    full_text = "".join(run.text for run in runs)
    start_idx = full_text.find(old_text)
    if start_idx == -1:
        return

    end_idx = start_idx + len(old_text)
    char_pos = 0
    placed = False

    for run in runs:
        rs = char_pos
        re_ = char_pos + len(run.text)

        if re_ <= start_idx or rs >= end_idx:
            pass
        elif rs >= start_idx and re_ <= end_idx:
            run.text = new_text if not placed else ""
            placed = True
        elif rs < start_idx < re_:
            prefix = run.text[: start_idx - rs]
            if re_ <= end_idx:
                run.text = prefix + new_text
                placed = True
            else:
                suffix = run.text[end_idx - rs :]
                run.text = prefix + new_text + suffix
                placed = True
        elif rs < end_idx < re_:
            suffix = run.text[end_idx - rs :]
            run.text = (new_text + suffix) if not placed else suffix
            placed = True

        char_pos = re_
