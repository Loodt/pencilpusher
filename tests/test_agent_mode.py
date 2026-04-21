"""Tests for agent-driven mode commands (read, detect, write-wiki, fill --field-map).

These commands are designed for use by external LLM agents (Claude Code, Codex)
that do the reasoning themselves — no Anthropic API calls needed.
"""

import json
from pathlib import Path

import pytest
from click.testing import CliRunner

from pencilpusher.cli import main


@pytest.fixture
def runner():
    return CliRunner()


@pytest.fixture
def vault_dir(tmp_path, monkeypatch):
    """Point pencilpusher at a temp vault and initialize it."""
    monkeypatch.setenv("PENCILPUSHER_VAULT", str(tmp_path / "vault"))
    result = CliRunner().invoke(main, ["init"])
    assert result.exit_code == 0
    return tmp_path / "vault"


class TestRead:
    def test_read_docx(self, runner, tmp_path):
        """Read a DOCX file and get markdown back."""
        from docx import Document

        doc = Document()
        doc.add_paragraph("Test Document Title")
        doc.add_paragraph("This is the body text.")
        docx_path = tmp_path / "test.docx"
        doc.save(str(docx_path))

        result = runner.invoke(main, ["read", str(docx_path)])
        assert result.exit_code == 0
        assert "Test Document Title" in result.output

    def test_read_nonexistent_file(self, runner):
        """Reading a nonexistent file should fail."""
        result = runner.invoke(main, ["read", "/nonexistent/file.pdf"])
        assert result.exit_code != 0


class TestDetect:
    def test_detect_acroform_pdf(self, runner):
        """Detect fields in an AcroForm PDF."""
        acroform_path = Path(__file__).parent / "test_form_acroform.pdf"
        if not acroform_path.exists():
            pytest.skip("test_form_acroform.pdf not available")

        result = runner.invoke(main, ["detect", str(acroform_path)])
        assert result.exit_code == 0

        data = json.loads(result.output)
        assert "fields" in data
        assert len(data["fields"]) > 0
        assert data["fields"][0]["field_type"] == "acroform"

    def test_detect_flat_pdf_warns(self, runner):
        """Flat PDFs should return empty fields with a warning."""
        flat_path = Path(__file__).parent / "test_form_flat.pdf"
        if not flat_path.exists():
            pytest.skip("test_form_flat.pdf not available")

        result = runner.invoke(main, ["detect", str(flat_path)])
        assert result.exit_code == 0

        data = json.loads(result.output)
        assert data["fields"] == []
        assert data["warning"] == "flat_pdf_requires_vision"

    def test_detect_docx_placeholders(self, runner, tmp_path):
        """Detect placeholder fields in a DOCX."""
        from docx import Document

        doc = Document()
        doc.add_paragraph("Full Name: [_______________]")
        doc.add_paragraph("ID Number: [_______________]")
        docx_path = tmp_path / "test.docx"
        doc.save(str(docx_path))

        result = runner.invoke(main, ["detect", str(docx_path)])
        assert result.exit_code == 0

        data = json.loads(result.output)
        assert "fields" in data
        assert len(data["fields"]) >= 2


class TestWriteWiki:
    def test_write_wiki_page(self, runner, vault_dir):
        """Write content to a wiki page and read it back."""
        result = runner.invoke(main, [
            "write-wiki", "identity",
            "# Identity\\nName: Jane Moyo\\nDOB: 1990-03-15",
        ])
        assert result.exit_code == 0
        assert "Updated wiki/identity.md" in result.output

        # Verify content was written
        content = (vault_dir / "wiki" / "identity.md").read_text(encoding="utf-8")
        assert "Jane Moyo" in content
        assert "1990-03-15" in content

    def test_write_wiki_stdin(self, runner, vault_dir):
        """Write wiki page from stdin."""
        result = runner.invoke(main, [
            "write-wiki", "banking", "--stdin",
        ], input="# Banking\nAccount: 123456\n")
        assert result.exit_code == 0
        assert "Updated wiki/banking.md" in result.output

        content = (vault_dir / "wiki" / "banking.md").read_text(encoding="utf-8")
        assert "123456" in content

    def test_write_wiki_company_page(self, runner, vault_dir):
        """Write to a company sub-page."""
        result = runner.invoke(main, [
            "write-wiki", "companies/acme",
            "# Acme Corp\\nReg: 123",
        ])
        assert result.exit_code == 0

        content = (vault_dir / "wiki" / "companies" / "acme.md").read_text(encoding="utf-8")
        assert "Acme Corp" in content

    def test_write_wiki_invalid_page(self, runner, vault_dir):
        """Writing to an unknown page name should fail."""
        result = runner.invoke(main, ["write-wiki", "nonsense_page", "content"])
        assert result.exit_code != 0

    def test_write_wiki_no_content(self, runner, vault_dir):
        """Writing with no content and no --stdin should fail."""
        result = runner.invoke(main, ["write-wiki", "identity"])
        assert result.exit_code != 0


class TestFillWithFieldMap:
    def test_fill_acroform_with_field_map(self, runner, tmp_path):
        """Fill an AcroForm PDF using explicit field mapping."""
        acroform_path = Path(__file__).parent / "test_form_acroform.pdf"
        if not acroform_path.exists():
            pytest.skip("test_form_acroform.pdf not available")

        # First detect to get field names
        detect_result = runner.invoke(main, ["detect", str(acroform_path)])
        data = json.loads(detect_result.output)
        if not data["fields"]:
            pytest.skip("No fields detected in test PDF")

        # Use the first field's name
        field_name = data["fields"][0]["field_key"]
        field_map = json.dumps({field_name: "Test Value"})

        output_path = tmp_path / "filled.pdf"
        result = runner.invoke(main, [
            "fill", str(acroform_path),
            "--field-map", field_map,
            "-o", str(output_path),
        ])
        assert result.exit_code == 0
        assert output_path.exists()

    def test_fill_docx_with_field_map(self, runner, tmp_path):
        """Fill a DOCX with placeholder fields using explicit mapping."""
        from docx import Document

        doc = Document()
        doc.add_paragraph("Full Name: [_______________]")
        doc.add_paragraph("ID Number: [_______________]")
        docx_path = tmp_path / "test.docx"
        doc.save(str(docx_path))

        field_map = json.dumps({
            "[_______________]": "Jane Moyo",
        })

        output_path = tmp_path / "filled.docx"
        result = runner.invoke(main, [
            "fill", str(docx_path),
            "--field-map", field_map,
            "-o", str(output_path),
        ])
        assert result.exit_code == 0
        assert output_path.exists()

    def test_fill_flat_pdf_with_fields_json(self, runner, tmp_path):
        """Fill a flat PDF using --field-map + --fields-json for positioning."""
        flat_path = Path(__file__).parent / "test_form_flat.pdf"
        if not flat_path.exists():
            pytest.skip("test_form_flat.pdf not available")

        # Agent provides both the values and the field positions
        field_map = json.dumps({"Full Name": "Jane Moyo"})
        fields_json = json.dumps([{
            "name": "Full Name",
            "field_type": "visual",
            "page": 0,
            "bbox": [15, 20, 50, 3],
        }])

        output_path = tmp_path / "filled_flat.pdf"
        result = runner.invoke(main, [
            "fill", str(flat_path),
            "--field-map", field_map,
            "--fields-json", fields_json,
            "-o", str(output_path),
        ])
        assert result.exit_code == 0
        assert output_path.exists()
        # Verify the output is larger than the input (widget added)
        assert output_path.stat().st_size > flat_path.stat().st_size

    def test_fill_invalid_json(self, runner, tmp_path):
        """Invalid JSON in --field-map should fail gracefully."""
        from docx import Document

        doc = Document()
        doc.add_paragraph("Name: [___]")
        docx_path = tmp_path / "test.docx"
        doc.save(str(docx_path))

        result = runner.invoke(main, [
            "fill", str(docx_path),
            "--field-map", "not-json",
        ])
        assert result.exit_code != 0

    def test_fill_invalid_fields_json(self, runner, tmp_path):
        """Invalid JSON in --fields-json should fail gracefully."""
        flat_path = Path(__file__).parent / "test_form_flat.pdf"
        if not flat_path.exists():
            pytest.skip("test_form_flat.pdf not available")

        result = runner.invoke(main, [
            "fill", str(flat_path),
            "--field-map", '{"Name": "Test"}',
            "--fields-json", "not-json",
        ])
        assert result.exit_code != 0


class TestProbe:
    """Tests for the `probe` command — flat-PDF layout introspection."""

    @staticmethod
    def _tiny_gridded_pdf(path: Path) -> Path:
        """Build a 1-page PDF with a 2x3 cell grid + digits '1'/'2'/'3' in the left col.

        Each vertical divider is drawn as 3 short segments (one per row) so
        the divider x-coord appears at least 3 times in `get_drawings()`,
        matching the default min_divider_count=3.
        """
        import fitz

        doc = fitz.open()
        page = doc.new_page(width=300, height=400)
        # Row boundaries at y=50/100/150/200 — each as 3 segments (one per column)
        # so each horizontal y also appears multiple times.
        row_ys = (50, 100, 150, 200)
        vert_xs = (30, 80, 270)
        for y in row_ys:
            for x0, x1 in zip(vert_xs[:-1], vert_xs[1:]):
                page.draw_line(fitz.Point(x0, y), fitz.Point(x1, y),
                               color=(0, 0, 0), width=0.4)
        # Column dividers — drawn as 3 segments (one per row) so the x-coord
        # accumulates count=3.
        for x in vert_xs:
            for y0, y1 in zip(row_ys[:-1], row_ys[1:]):
                page.draw_line(fitz.Point(x, y0), fitz.Point(x, y1),
                               color=(0, 0, 0), width=0.4)
        # Row numbers '1', '2', '3' inside the No. column (x ~ 40-65).
        for row_num, y in enumerate((68, 118, 168), start=1):
            page.insert_text(fitz.Point(50, y), str(row_num),
                             fontname="helv", fontsize=10)
        doc.save(str(path))
        doc.close()
        return path

    def test_probe_reports_column_dividers_and_digits(self, runner, tmp_path):
        """probe should return column dividers and digit spans for the grid."""
        pdf = self._tiny_gridded_pdf(tmp_path / "grid.pdf")

        result = runner.invoke(main, ["probe", str(pdf)])
        assert result.exit_code == 0, result.output

        data = json.loads(result.output)
        assert "pages" in data and len(data["pages"]) == 1

        page = data["pages"][0]
        assert page["page"] == 0
        assert page["width"] == 300.0
        assert page["height"] == 400.0

        # With min-divider-count=3 (default), the three vertical lines
        # at x=30/80/270 should all be reported (they each appear in at
        # least 3 drawn segments via their endpoints).
        dividers = page["column_dividers"]
        # Expect at least one divider near each of 30, 80, 270 (tolerance 2pt).
        for target in (30, 80, 270):
            assert any(abs(d - target) < 2 for d in dividers), (
                f"expected column divider near x={target}, got {dividers}"
            )

        # Three digit spans "1", "2", "3" should appear.
        digit_texts = sorted(d["text"] for d in page["digit_spans"])
        assert digit_texts == ["1", "2", "3"]
        # Their x0 should sit in the No.-column range (30-80).
        for d in page["digit_spans"]:
            x0 = d["bbox"][0]
            assert 30 <= x0 <= 80, d

    def test_probe_rejects_non_pdf(self, runner, tmp_path):
        """probe should reject non-PDF inputs with a non-zero exit."""
        from docx import Document

        doc = Document()
        doc.add_paragraph("Not a PDF")
        docx_path = tmp_path / "not-a-pdf.docx"
        doc.save(str(docx_path))

        result = runner.invoke(main, ["probe", str(docx_path)])
        assert result.exit_code != 0


class TestFillTextboxMode:
    """Tests for `fill --textbox-mode` on flat PDFs."""

    @staticmethod
    def _blank_flat_pdf(path: Path) -> Path:
        import fitz

        doc = fitz.open()
        doc.new_page(width=595, height=842)
        doc.save(str(path))
        doc.close()
        return path

    def test_textbox_mode_wraps_long_answer(self, runner, tmp_path):
        """A long answer should fit inside a narrow cell via font-shrink."""
        import fitz

        pdf = self._blank_flat_pdf(tmp_path / "flat.pdf")

        # 100 x 20 pt cell at (40, 40) — too small for a 120-char answer
        # at 10 pt, so the shrink fallback must kick in.
        long_value = ("Spillage water from Elution/Carbon bunded sump; "
                      "alkaline pH 10-11; residual CN- ~50-100 mg/L")
        # Convert to percentages of the 595 x 842 page.
        bbox = [40 / 595 * 100, 40 / 842 * 100,
                100 / 595 * 100, 20 / 842 * 100]
        fields_json = json.dumps([{
            "name": "Stream",
            "field_type": "visual",
            "page": 0,
            "bbox": bbox,
            "textbox_options": {
                "font_size": 10,
                "font_color": [0, 0, 0.75],
            },
        }])
        field_map = json.dumps({"Stream": long_value})

        output_path = tmp_path / "filled.pdf"
        result = runner.invoke(main, [
            "fill", str(pdf),
            "--field-map", field_map,
            "--fields-json", fields_json,
            "--textbox-mode",
            "-o", str(output_path),
        ])
        assert result.exit_code == 0, result.output
        assert output_path.exists()

        # In textbox mode we should NOT have added an AcroForm widget;
        # text goes straight onto the page as content.
        out_doc = fitz.open(str(output_path))
        try:
            page = out_doc[0]
            assert len(list(page.widgets())) == 0
            # Some portion of the answer must be rendered as text.
            text = page.get_text()
            # At the minimum, a recognisable fragment should appear — the
            # shrink fallback may truncate, so we check a short prefix.
            assert "Spill" in text or "water" in text, text
        finally:
            out_doc.close()

    def test_widget_mode_still_default(self, runner, tmp_path):
        """Without --textbox-mode, flat-PDF fill still creates a widget."""
        import fitz

        pdf = self._blank_flat_pdf(tmp_path / "flat.pdf")
        fields_json = json.dumps([{
            "name": "Name",
            "field_type": "visual",
            "page": 0,
            "bbox": [10, 10, 30, 3],
        }])
        field_map = json.dumps({"Name": "Jane"})

        output_path = tmp_path / "filled.pdf"
        result = runner.invoke(main, [
            "fill", str(pdf),
            "--field-map", field_map,
            "--fields-json", fields_json,
            "-o", str(output_path),
        ])
        assert result.exit_code == 0, result.output

        out_doc = fitz.open(str(output_path))
        try:
            widgets = list(out_doc[0].widgets())
            assert len(widgets) == 1
            assert widgets[0].field_value == "Jane"
        finally:
            out_doc.close()
