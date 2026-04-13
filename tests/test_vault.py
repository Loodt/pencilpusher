"""Tests for the vault store and crypto modules."""


import pytest

from pencilpusher.vault.crypto import decrypt_text, encrypt_text, init_vault_encryption
from pencilpusher.vault.store import Vault


class TestCrypto:
    def test_encrypt_decrypt_roundtrip(self, tmp_path):
        fernet = init_vault_encryption(tmp_path, "test-password")
        plaintext = "My secret ID number: 8501015800088"
        encrypted = encrypt_text(fernet, plaintext)
        decrypted = decrypt_text(fernet, encrypted)
        assert decrypted == plaintext

    def test_wrong_password_raises(self, tmp_path):
        init_vault_encryption(tmp_path, "correct-password")
        with pytest.raises(ValueError, match="Wrong password"):
            init_vault_encryption(tmp_path, "wrong-password")

    def test_salt_persists(self, tmp_path):
        fernet1 = init_vault_encryption(tmp_path, "my-password")
        encrypted = encrypt_text(fernet1, "hello")

        # Re-open with same password should work
        fernet2 = init_vault_encryption(tmp_path, "my-password")
        assert decrypt_text(fernet2, encrypted) == "hello"


class TestVault:
    def test_init_creates_structure(self, tmp_path):
        vault = Vault(vault_dir=tmp_path, password=None)
        vault.init()

        assert (tmp_path / "raw").is_dir()
        assert (tmp_path / "wiki").is_dir()
        assert (tmp_path / "output").is_dir()
        assert (tmp_path / "wiki" / "index.md").exists()
        assert (tmp_path / "wiki" / "log.md").exists()
        assert (tmp_path / "wiki" / "identity.md").exists()

    def test_wiki_page_roundtrip(self, tmp_path):
        vault = Vault(vault_dir=tmp_path, password=None)
        vault.init()

        vault.write_wiki_page("identity", "# Identity\n\n**Name:** Test User\n")
        content = vault.read_wiki_page("identity")
        assert "Test User" in content

    def test_read_all_wiki_pages(self, tmp_path):
        vault = Vault(vault_dir=tmp_path, password=None)
        vault.init()

        pages = vault.read_all_wiki_pages()
        assert "identity" in pages
        assert "banking" in pages

    def test_store_raw_file(self, tmp_path):
        vault = Vault(vault_dir=tmp_path, password=None)
        vault.init()

        # Create a fake source file
        source = tmp_path / "test_doc.pdf"
        source.write_text("fake pdf content")

        stored = vault.store_raw(source, "identity")
        assert stored.exists()

        files = vault.list_raw_files()
        assert len(files) == 1
        assert files[0]["category"] == "identity"

    def test_vault_with_password(self, tmp_path):
        """Vault with password still works (encryption optional, requires cryptography pkg)."""
        try:
            vault = Vault(vault_dir=tmp_path, password="test123")
        except ImportError:
            pytest.skip("cryptography package not installed")
        vault.init()

        vault.write_wiki_page("identity", "# Identity\n\n**ID:** 8501015800088\n")
        content = vault.read_wiki_page("identity")
        assert "8501015800088" in content

        # Plaintext file should exist (MVP: no .enc files)
        md_path = tmp_path / "wiki" / "identity.md"
        assert md_path.exists()


class TestDetector:
    """Basic tests for field detection (no LLM needed)."""

    def test_detect_docx_placeholder_brackets(self, tmp_path):
        """Test detection of [___] style placeholders."""
        from docx import Document

        from pencilpusher.fill.detector import detect_docx_fields

        doc = Document()
        doc.add_paragraph("Full Name: [_______________]")
        doc.add_paragraph("ID Number: [_______________]")
        docx_path = tmp_path / "test.docx"
        doc.save(str(docx_path))

        fields = detect_docx_fields(docx_path)
        assert len(fields) >= 2

    def test_detect_docx_placeholder_underlines(self, tmp_path):
        """Test detection of _________ style placeholders."""
        from docx import Document

        from pencilpusher.fill.detector import detect_docx_fields

        doc = Document()
        doc.add_paragraph("Name: ___________________")
        docx_path = tmp_path / "test.docx"
        doc.save(str(docx_path))

        fields = detect_docx_fields(docx_path)
        assert len(fields) >= 1
