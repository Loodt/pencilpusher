"""pencilpusher demo — proves the full pipeline works with a fictional user.

This demo does NOT require an Anthropic API key. It manually builds the wiki
from the example source documents and then fills both a DOCX and PDF form.

Usage:
    python examples/run_demo.py

What it does:
    1. Creates a temporary vault
    2. Builds a wiki from Jane Moyo's fictional source documents
    3. Fills an application form (DOCX) using the wiki data
    4. Fills the same form (PDF) using the wiki data
    5. Shows the results
"""

import shutil
import sys
from pathlib import Path

# Add src to path for development
sys.path.insert(0, str(Path(__file__).parent.parent / "src"))

from docx import Document
from docx.shared import Pt

EXAMPLES_DIR = Path(__file__).parent
SOURCES_DIR = EXAMPLES_DIR / "sources"
INBOX_DIR = EXAMPLES_DIR / "inbox"
OUTPUT_DIR = EXAMPLES_DIR / "expected"


# === Jane Moyo's data (extracted from source documents) ===
# In production, Claude extracts this. For the demo, we hardcode it.

WIKI = {
    "identity": """# Identity

**Full Name:** JANE THANDEKA MOYO
**Date of Birth:** 15 March 1990
**Gender:** Female
**Nationality:** Zambian / South African (dual)
**NRC Number:** 445566/77/1
**SA ID Number:** 9003155141085
**Passport Number:** ZN0098765 (SA, expires 09 Jan 2034)
**Place of Birth:** Lusaka, Zambia

---
*Sources: jane_moyo_id.md, jane_moyo_passport.md*
""",
    "contacts": """# Contacts

**Phone:** +27 61 234 5678
**Email:** jane.moyo@example.com

---
*Source: jane_moyo_passport.md*
""",
    "addresses": """# Addresses

**Physical Address:** 42 Protea Lane, Sandton, Gauteng, 2196, South Africa

---
*Source: jane_moyo_passport.md*
""",
    "companies/moyo-mining": """# Moyo Mining Solutions Limited

**Company Name:** MOYO MINING SOLUTIONS LIMITED
**Registration Number:** 120240067890
**Date of Incorporation:** 12 January 2024
**Registered Office:** Plot 15, Makeni Road, Lusaka, Zambia
**Postal Address:** P.O. Box 31234, Lusaka, 10101, Zambia
**Phone:** +260 97 123 4567
**Email:** info@moyomining.example.com
**Nominal Capital:** K 50,000
**Share Capital:** 50,000 Ordinary shares at K1.00

## Directors
| Name | Nationality | NRC | Role | Shares |
|------|-------------|-----|------|--------|
| JANE THANDEKA MOYO | Zambian | 445566/77/1 | Managing Director + Secretary | 35,000 |
| DAVID BANDA | Zambian | 223344/55/1 | Operations Director | 15,000 |

## Beneficial Owners
- JANE THANDEKA MOYO — 35,000 shares (70%) — Direct
- DAVID BANDA — 15,000 shares (30%) — Direct

---
*Source: moyo_mining_pacra.md (PACRA printout 01/04/2026)*
""",
}

# Data for form filling (what the matcher would produce)
FILL_DATA = {
    "Full Name": "JANE THANDEKA MOYO",
    "Date of Birth": "15 March 1990",
    "Gender": "Female",
    "Nationality": "Zambian",
    "Identity Type": "NRC",
    "Identity Number": "445566/77/1",
    "Phone Number": "+27 61 234 5678",
    "Email Address": "jane.moyo@example.com",
    "Physical Address": "42 Protea Lane, Sandton, Gauteng, 2196, South Africa",
    "Company Name": "MOYO MINING SOLUTIONS LIMITED",
    "Registration Number": "120240067890",
    "Registered Office": "Plot 15, Makeni Road, Lusaka, Zambia",
    "Position/Title": "Managing Director",
}


def fill_docx_form():
    """Fill the example DOCX form with Jane Moyo's data."""
    doc = Document(str(INBOX_DIR / "application_form.docx"))
    table = doc.tables[0]

    for row in table.rows:
        label_text = row.cells[0].text.strip().rstrip(":")
        if label_text in FILL_DATA:
            cell = row.cells[1]
            if cell.paragraphs:
                run = cell.paragraphs[0].add_run(FILL_DATA[label_text])
                run.font.size = Pt(10)

    output = OUTPUT_DIR / "application_form_filled.docx"
    doc.save(str(output))
    return output


def fill_pdf_form():
    """Fill the example PDF form with Jane Moyo's data."""
    import fitz

    doc = fitz.open(str(INBOX_DIR / "application_form.pdf"))
    page = doc[0]

    # Extract text positions to find labels
    spans = []
    for block in page.get_text("dict")["blocks"]:
        if "lines" not in block:
            continue
        for line in block["lines"]:
            for span in line["spans"]:
                text = span["text"].strip()
                if text:
                    spans.append((text, span["bbox"], span.get("size", 10)))

    filled = 0
    for text, bbox, fontsize in spans:
        label = text.rstrip(":")
        if label in FILL_DATA and bbox[0] < 200:
            value = FILL_DATA[label]
            point = fitz.Point(222, bbox[1] + fontsize)
            page.insert_text(point, value, fontsize=fontsize, fontname="helv", color=(0, 0, 0))
            filled += 1

    output = OUTPUT_DIR / "application_form_filled.pdf"
    doc.save(str(output))
    doc.close()
    return output, filled


def main():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    print("=" * 60)
    print("  pencilpusher demo — Jane Moyo (fictional user)")
    print("=" * 60)

    # Step 1: Show source documents
    print("\n--- Source Documents ---")
    for f in sorted(SOURCES_DIR.glob("*.md")):
        print(f"  {f.name}")

    # Step 2: Show wiki (what Claude would build)
    print("\n--- Wiki Pages (built from sources) ---")
    for page_name, content in WIKI.items():
        lines = content.strip().split("\n")
        title = lines[0].lstrip("# ")
        field_count = content.count("**")
        print(f"  wiki/{page_name}.md — {title} ({field_count // 2} fields)")

    # Step 3: Show forms to fill
    print("\n--- Forms to Fill ---")
    for f in sorted(INBOX_DIR.iterdir()):
        if f.is_file():
            print(f"  inbox/{f.name} ({f.stat().st_size:,} bytes)")

    # Step 4: Fill DOCX
    print("\n--- Filling DOCX Form ---")
    docx_output = fill_docx_form()
    print(f"  Filled {len(FILL_DATA)} fields")
    print(f"  Output: {docx_output}")

    # Step 5: Fill PDF
    print("\n--- Filling PDF Form ---")
    pdf_output, pdf_count = fill_pdf_form()
    print(f"  Filled {pdf_count} fields")
    print(f"  Output: {pdf_output}")

    # Step 6: Summary
    print("\n--- Results ---")
    print(f"  DOCX: {docx_output.name} ({docx_output.stat().st_size:,} bytes)")
    print(f"  PDF:  {pdf_output.name} ({pdf_output.stat().st_size:,} bytes)")
    print()
    print("Open both files to verify:")
    print(f"  {docx_output}")
    print(f"  {pdf_output}")
    print()
    print("Both should show Jane Moyo's data in the correct fields")
    print("with original form styling preserved.")


if __name__ == "__main__":
    main()
